import argparse
import os
import re
import shutil
from langchain.schema.document import Document
from get_embedding_function import get_embedding_function
from langchain_chroma import Chroma

CHROMA_PATH = "chroma"
DATA_PATH = "data"
CHUNK_SIZE = 800
CHUNK_OVERLAP = 80


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--reset", action="store_true", help="Мэдээллийн санг устгаж дахин үүсгэх.")
    args = parser.parse_args()
    if args.reset:
        print("Мэдээллийн санг цэвэрлэж байна...")
        clear_database()

    documents = load_documents()
    if not documents:
        print("Уншигдсан баримт бичиг олдсонгүй. 'data/' хавтсанд файл байгаа эсэхийг шалгана уу.")
        return

    chunks = split_documents(documents)
    add_to_chroma(chunks)

def load_documents() -> list:
    documents = []

    for filename in sorted(os.listdir(DATA_PATH)):
        file_path = os.path.join(DATA_PATH, filename)

        if filename.lower().endswith(".pdf"):
            docs = _load_pdf(file_path, filename)
            documents.extend(docs)

        elif filename.lower().endswith(".docx"):
            docs = _load_docx(file_path, filename)
            documents.extend(docs)

        elif not filename.startswith("."):
            print(f"Дэмжигдээгүй формат алгасав: {filename}")

    print(f"\nНийт уншигдсан хэсэг: {len(documents)}")
    return documents


def _load_pdf(file_path: str, filename: str) -> list:
    # 1. pypdf
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(Document(
                    page_content=text,
                    metadata={"source": file_path, "page": i, "loader": "pypdf"},
                ))
        if pages:
            print(f"{filename}: pypdf-ээр {len(pages)} хуудас уншигдлаа.")
            return pages
    except Exception as e:
        print(f"{filename} pypdf алдаа: {e}")

    # 2. pdfplumber
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                for table in page.extract_tables():
                    for row in table:
                        text += "\n" + " | ".join(cell or "" for cell in row)
                text = text.strip()
                if text:
                    pages.append(Document(
                        page_content=text,
                        metadata={"source": file_path, "page": i, "loader": "pdfplumber"},
                    ))
        if pages:
            print(f"{filename}: pdfplumber-ээр {len(pages)} хуудас уншигдлаа.")
            return pages
    except Exception as e:
        print(f"{filename} pdfplumber алдаа: {e}")

    # 3. OCR
    try:
        from pdf2image import convert_from_path
        import pytesseract
        POPPLER_PATH = "/opt/homebrew/opt/poppler/bin"
        print(f"{filename}: Зурган PDF, OCR эхэлж байна...")
        images = convert_from_path(file_path, dpi=200, poppler_path=POPPLER_PATH)
        pages = []
        for i, img in enumerate(images):
            try:
                text = pytesseract.image_to_string(img, lang="mon+eng").strip()
            except Exception:
                text = pytesseract.image_to_string(img, lang="eng").strip()
            if text:
                pages.append(Document(
                    page_content=text,
                    metadata={"source": file_path, "page": i, "loader": "ocr"},
                ))
        if pages:
            print(f"{filename}: OCR-ээр {len(pages)} хуудас уншигдлаа.")
            return pages
    except ImportError:
        print("OCR: pip install pdf2image pytesseract")
    except Exception as e:
        print(f"{filename} OCR алдаа: {e}")

    return []


def _load_docx(file_path: str, filename: str) -> list:
    # 1. python-docx
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)
        full_text = "\n".join(parts).strip()
        if full_text:
            print(f"{filename}: python-docx-ээр уншигдлаа ({len(full_text)} тэмдэгт).")
            return [Document(
                page_content=full_text,
                metadata={"source": file_path, "page": 0, "loader": "python-docx"},
            )]
    except Exception as e:
        print(f"{filename} python-docx алдаа: {e}")

    # 2. docx2txt backup
    try:
        import docx2txt
        text = docx2txt.process(file_path).strip()
        if text:
            print(f"{filename}: docx2txt-ээр уншигдлаа.")
            return [Document(
                page_content=text,
                metadata={"source": file_path, "page": 0, "loader": "docx2txt"},
            )]
    except Exception as e:
        print(f"{filename} docx2txt алдаа: {e}")

    return []

MN_SENTENCE_END = re.compile(
    r'('
    r'оршино\.|байна\.|болно\.|юм\.|байжээ\.|болжээ\.|гэжээ\.'
    r'|байв\.|болов\.|гэв\.'
    r'|сан\.|сэн\.|сон\.|сөн\.'
    r'|лаа\.|лээ\.|лоо\.|лөө\.'
    r'|дэнэ\.|дэнэ\.'
    r'|на\.|нэ\.|но\.|нө\.'
    r'|мөрдөнө\.|хамаарахгүй\.|зохицуулна\.'
    r'|[!?]'
    r')',
    re.UNICODE
)
HEADING_PATTERN = re.compile(
    r'[А-ЯӨҮЁ]{3,}(?:\s+[А-ЯӨҮЁ]+)*' 
    r')',
    re.UNICODE
)

SUB_NUMBER = re.compile(r'^\d+\.\d+\.')


def split_text_by_sentences(text: str) -> list:
    parts = MN_SENTENCE_END.split(text)
    sentences = []
    i = 0
    while i < len(parts):
        chunk = parts[i].strip()
        if i + 1 < len(parts) and MN_SENTENCE_END.match(parts[i + 1]):
            chunk = chunk + parts[i + 1]
            i += 2
        else:
            i += 1
        if chunk:
            sentences.append(chunk)
    return sentences

def split_into_paragraphs(text: str) -> list:
    text = re.sub(r'\nХэвлэх\n|\nХэвлэгдсэн\n', '\n', text)
    text = re.sub(r'\nХэвлэх$', '', text)
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    paragraphs = []
    current_heading = ""
    current_content = []
    for line in lines:
        if SUB_NUMBER.match(line):
            if current_content:
                para = (current_heading + " " + " ".join(current_content)).strip()
                paragraphs.append(para)
                current_heading = ""
                current_content = []
            current_content = [line]
        elif HEADING_PATTERN.match(line) and not SUB_NUMBER.match(line):
            if current_content:
                para = (current_heading + " " + " ".join(current_content)).strip()
                paragraphs.append(para)
                current_content = []
            elif current_heading:
                paragraphs.append(current_heading)
            current_heading = line
        else:
            current_content.append(line)
    if current_content:
        para = (current_heading + " " + " ".join(current_content)).strip()
        paragraphs.append(para)
    elif current_heading:
        paragraphs.append(current_heading)
    return [p for p in paragraphs if p]


def split_documents(documents: list) -> list:

    all_chunks = []

    for doc in documents:
        paragraphs = split_into_paragraphs(doc.page_content)
        for para in paragraphs:
            if len(para) <= CHUNK_SIZE:
                all_chunks.append(Document(
                    page_content=para,
                    metadata=dict(doc.metadata),
                ))
            else:
                sentences = split_text_by_sentences(para)
                sub_chunks = _build_chunks(sentences, doc.metadata)
                all_chunks.extend(sub_chunks)

    print(f"Нийт {len(all_chunks)} chunk үүсгэгдлээ.")
    return all_chunks


def _build_chunks(sentences: list, metadata: dict) -> list:
    chunks = []
    current_sentences = []
    current_len = 0

    for sentence in sentences:
        sentence_len = len(sentence)

        if current_len + sentence_len > CHUNK_SIZE and current_sentences:
            chunk_text = " ".join(current_sentences)
            chunks.append(Document(
                page_content=chunk_text,
                metadata=dict(metadata),
            ))

            overlap_sentences = []
            overlap_len = 0
            for s in reversed(current_sentences):
                if overlap_len + len(s) <= CHUNK_OVERLAP:
                    overlap_sentences.insert(0, s)
                    overlap_len += len(s)
                else:
                    break

            current_sentences = overlap_sentences
            current_len = overlap_len

        current_sentences.append(sentence)
        current_len += sentence_len

    if current_sentences:
        chunk_text = " ".join(current_sentences)
        chunks.append(Document(
            page_content=chunk_text,
            metadata=dict(metadata),
        ))

    return chunks

def add_to_chroma(chunks: list):
    db = Chroma(
        persist_directory=CHROMA_PATH,
        embedding_function=get_embedding_function(),
    )
    chunks_with_ids = calculate_chunk_ids(chunks)

    existing_items = db.get(include=[])
    existing_ids = set(existing_items["ids"])
    print(f"Мэдээллийн санд одоо байгаа chunk: {len(existing_ids)}")

    new_chunks = [c for c in chunks_with_ids if c.metadata["id"] not in existing_ids]

    if new_chunks:
        print(f"Шинээр нэмэгдэх chunk: {len(new_chunks)}")
        new_chunk_ids = [c.metadata["id"] for c in new_chunks]
        db.add_documents(new_chunks, ids=new_chunk_ids)
    else:
        print("Шинэ chunk алга. Мэдээллийн сан хэвийн байна.")


def calculate_chunk_ids(chunks: list) -> list:
    last_page_id = None
    current_chunk_index = 0

    for chunk in chunks:
        source = chunk.metadata.get("source", "unknown")
        page = chunk.metadata.get("page", 0)
        current_page_id = f"{source}:{page}"

        if current_page_id == last_page_id:
            current_chunk_index += 1
        else:
            current_chunk_index = 0

        chunk.metadata["id"] = f"{current_page_id}:{current_chunk_index}"
        last_page_id = current_page_id

    return chunks

def clear_database():
    if os.path.exists(CHROMA_PATH):
        shutil.rmtree(CHROMA_PATH)
        print("Мэдээллийн сан устгагдлаа.")


if __name__ == "__main__":
    main()